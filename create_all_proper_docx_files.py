#!/usr/bin/env python3
"""
Script to create all blog posts and service documents as proper Word documents
"""

import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Complete list of all services
ALL_SERVICES = [
    "house-cleaning",
    "move-in-move-out-cleaning", 
    "pressure-washing",
    "dog-walking",
    "graphic-design",
    "business-coaching",
    "virtual-assistant",
    "gutter-cleaning",
    "window-cleaning",
    "carpet-shampooing",
    "garage-attic-organization",
    "trash-hauling-junk-removal",
    "airbnb-cleaning",
    "lawn-mowing-yard-maintenance",
    "furniture-assembly",
    "tv-mounting",
    "handyman-services",
    "fence-painting",
    "light-bulb-fixture-installation",
    "basic-drywall-patching",
    "mailbox-installation",
    "holiday-light-hanging",
    "lockout-assistance",
    "pool-cleaning",
    "grocery-shopping-delivery",
    "prescription-pickup",
    "waiting-in-line-service",
    "personal-assistant-service",
    "moving-assistance",
    "courier-delivery-services",
    "plant-watering",
    "decluttering-service",
    "packing-unpacking-service",
    "dog-waste-cleanup",
    "babysitting",
    "pet-sitting",
    "pet-taxi",
    "mobile-pet-grooming",
    "pet-poop-scooping-service",
    "pet-yard-deodorizing",
    "aquarium-cleaning",
    "home-safety-baby-proofing",
    "parent-helper-mothers-helper",
    "toy-organization-service",
    "birthday-party-setup-hosting",
    "party-rental-setup",
    "photography",
    "videography-for-events",
    "logo-design",
    "tshirt-merch-design",
    "print-on-demand-order-management",
    "product-sourcing",
    "drop-off-ebay-amazon-seller",
    "furniture-flipping",
    "home-office-gaming-setup-installer",
    "bookkeeping",
    "data-entry",
    "email-inbox-management",
    "crm-data-organization-setup",
    "transcription-services",
    "content-writing-blogging",
    "social-media-management",
    "social-media-consulting",
    "marketing-consulting",
    "customer-service-outsourcing",
    "cold-calling-appointment-setting",
    "online-research-assistant",
    "voiceover-work",
    "resume-writing",
    "life-coaching",
    "confidence-public-speaking-coaching",
    "relationship-coaching",
    "productivity-time-management-coaching",
    "accountability-coaching"
]

# General blog post topics
BLOG_POSTS = [
    "10-ways-to-organize-your-home-office",
    "creating-a-sustainable-home",
    "decluttering-your-home-room-by-room",
    "eco-friendly-cleaning-tips",
    "essential-home-safety-tips",
    "remote-work-productivity-tips",
    "seasonal-home-maintenance-checklist",
    "starting-a-home-service-business",
    "time-management-tips-for-busy-professionals",
    "maximizing-small-spaces",
    "smart-home-technology-guide",
    "home-renovation-budget-tips",
    "creating-a-peaceful-home-environment",
    "energy-efficient-home-improvements",
    "home-organization-for-families",
    "minimalist-living-tips",
    "home-security-best-practices",
    "creating-a-home-gym",
    "home-office-ergonomics",
    "sustainable-gardening-tips"
]

def create_service_blog_content(service_name):
    """Create comprehensive content for a service blog post"""
    title = service_name.replace('-', ' ').replace('_', ' ').title()
    
    content = f"""# {title}

## Meta Title: {title} - Professional {title} Services

## Meta Description: Discover professional {title.lower()} services. Get expert help with {title.lower()} for your home or business needs.

Content:

## What is {title}?

{title} is a specialized service that provides professional assistance for {title.lower()} needs. Whether you're a busy homeowner, a small business owner, or someone looking for expert help, our {title.lower()} services are designed to meet your specific requirements.

## Why Choose Professional {title} Services?

Professional {title.lower()} services offer several advantages:

- **Expert Knowledge**: Trained professionals with years of experience
- **Quality Results**: Consistent, high-quality work every time
- **Time Savings**: Free up your valuable time for other priorities
- **Peace of Mind**: Reliable, insured, and guaranteed services
- **Customized Solutions**: Tailored to your specific needs and preferences

## Our {title} Process

### 1. Initial Consultation
We start with a thorough consultation to understand your specific needs and requirements.

### 2. Customized Plan
Based on your consultation, we create a customized plan that addresses your unique situation.

### 3. Professional Execution
Our trained professionals execute the plan with attention to detail and quality standards.

### 4. Quality Assurance
We ensure all work meets our high standards before considering the job complete.

### 5. Follow-up Support
We provide ongoing support and are available for any questions or additional needs.

## Service Features

- **Flexible Scheduling**: Services available at times that work for you
- **Competitive Pricing**: Fair, transparent pricing with no hidden fees
- **Professional Equipment**: State-of-the-art tools and materials
- **Insurance Coverage**: Fully insured for your peace of mind
- **Satisfaction Guarantee**: We're not satisfied until you are

## Pricing Information

Our {title.lower()} services are competitively priced and offer excellent value:

- **Basic Package**: Starting from $XX
- **Standard Package**: Starting from $XX
- **Premium Package**: Starting from $XX

*Pricing varies based on scope, location, and specific requirements. Contact us for a personalized quote.*

## Service Areas

We provide {title.lower()} services in the following areas:
- Residential homes
- Commercial properties
- Small businesses
- Special events and occasions

## Customer Testimonials

*"The {title.lower()} service exceeded my expectations. Professional, reliable, and excellent quality work."* - Sarah M.

*"I highly recommend their {title.lower()} services. They were punctual, professional, and did an amazing job."* - John D.

*"Outstanding {title.lower()} service! They transformed my space and I couldn't be happier."* - Lisa R.

## How to Get Started

Getting started with our {title.lower()} services is easy:

1. **Contact Us**: Reach out via phone, email, or our online form
2. **Schedule Consultation**: We'll arrange a convenient time to discuss your needs
3. **Receive Quote**: Get a detailed, no-obligation quote
4. **Book Service**: Schedule your service at a time that works for you
5. **Enjoy Results**: Sit back and enjoy the professional results

## Contact Information

- **Phone**: [Your Phone Number]
- **Email**: [Your Email]
- **Website**: [Your Website]
- **Service Hours**: Monday - Friday: 8 AM - 6 PM, Saturday: 9 AM - 4 PM

## Frequently Asked Questions

**Q: How long does a typical {title.lower()} service take?**
A: Service duration varies based on scope and requirements. We'll provide a detailed timeline during consultation.

**Q: Do you offer emergency {title.lower()} services?**
A: Yes, we offer emergency services for urgent situations. Contact us immediately for emergency assistance.

**Q: Are your {title.lower()} services insured?**
A: Yes, all our services are fully insured for your protection and peace of mind.

**Q: Do you provide {title.lower()} services on weekends?**
A: Yes, we offer flexible scheduling including weekend appointments for your convenience.

**Q: What areas do you serve for {title.lower()}?**
A: We serve [Your Service Area]. Contact us to confirm availability in your specific location.

## Why Choose Us for {title}?

- **Experience**: Years of professional experience in {title.lower()}
- **Reliability**: Consistent, dependable service you can count on
- **Quality**: We never compromise on quality or customer satisfaction
- **Value**: Competitive pricing for premium service
- **Support**: Ongoing support and assistance when you need it

## Get Your Free Quote Today

Don't wait to experience the difference professional {title.lower()} services can make. Contact us today for your free quote!

**Call us now**: [Your Phone Number]
**Email us**: [Your Email]
**Visit our website**: [Your Website]

*Professional {title.lower()} services - when quality and reliability matter.*

---

*This blog post is designed to provide comprehensive information about {title.lower()} services while incorporating SEO best practices and engaging content for potential customers.*"""
    return content

def create_general_blog_content(topic):
    """Create content for general blog posts"""
    title = topic.replace('-', ' ').replace('_', ' ').title()
    
    content = f"""# {title}

## Meta Title: {title} - Complete Guide and Tips

## Meta Description: Learn everything about {title.lower()} with our comprehensive guide. Get expert tips and professional advice.

Content:

## Introduction

{title} is an essential aspect of modern living that can significantly improve your quality of life. Whether you're a homeowner, business owner, or simply looking to enhance your daily routine, understanding {title.lower()} can help you create better environments and more efficient systems.

## Why {title} is Important

Understanding and implementing {title.lower()} provides numerous benefits:

- **Improved Efficiency**: Streamlined processes save time and energy
- **Better Organization**: Systematic approaches reduce stress and confusion
- **Enhanced Safety**: Proper practices protect you and your family
- **Cost Savings**: Efficient systems reduce unnecessary expenses
- **Quality of Life**: Better environments improve overall well-being

## Key Components of {title}

### Planning and Preparation
- Assess your current situation
- Identify areas for improvement
- Set realistic goals and timelines
- Gather necessary resources and tools
- Create detailed action plans

### Implementation Strategies
- Start with high-impact changes
- Build sustainable habits gradually
- Monitor progress and adjust as needed
- Celebrate small victories
- Maintain momentum through consistency

### Maintenance and Optimization
- Regular reviews and assessments
- Continuous improvement processes
- Adaptation to changing needs
- Long-term sustainability planning
- Professional support when needed

## Professional Services

Our expert team can help you with {title.lower()}:

- **Consultation Services**: Professional assessment and planning
- **Implementation Support**: Hands-on assistance with changes
- **Training and Education**: Learn best practices and techniques
- **Ongoing Maintenance**: Regular check-ins and adjustments
- **Custom Solutions**: Tailored approaches for your specific needs

## Success Stories

*"Professional {title.lower()} services transformed our home and daily routine."* - Sarah M.

*"The expert guidance helped us create systems that actually work for our family."* - John D.

*"We never realized how much difference proper {title.lower()} could make."* - Lisa R.

## Getting Started

Ready to improve your {title.lower()}?

1. **Contact Us**: Schedule a consultation
2. **Assessment**: We'll evaluate your current situation
3. **Custom Plan**: Receive a personalized strategy
4. **Implementation**: Professional setup and training
5. **Ongoing Support**: Regular maintenance and improvements

## Contact Information

- **Phone**: [Your Phone Number]
- **Email**: [Your Email]
- **Website**: [Your Website]
- **Service Hours**: Monday - Friday: 8 AM - 6 PM, Saturday: 9 AM - 4 PM

---

*Transform your life with professional {title.lower()} services.*"""
    return content

def add_markdown_paragraph(doc, text):
    """Parse markdown-style text and add to docx document with formatting."""
    # Handle headings
    heading_match = re.match(r'^(#+) (.*)', text)
    if heading_match:
        level = len(heading_match.group(1))
        heading_text = heading_match.group(2).strip()
        doc.add_heading(heading_text, min(level, 4))
        return
    # Handle lists (bulleted or numbered)
    list_match = re.match(r'^[-*] (.*)', text)
    if list_match:
        # Add as a bulleted list item
        para = doc.add_paragraph(style='List Bullet')
        add_runs_with_formatting(para, list_match.group(1))
        return
    num_match = re.match(r'^(\d+)\. (.*)', text)
    if num_match:
        para = doc.add_paragraph(style='List Number')
        add_runs_with_formatting(para, num_match.group(2))
        return
    # Otherwise, treat as a normal paragraph
    para = doc.add_paragraph()
    add_runs_with_formatting(para, text)

def add_runs_with_formatting(para, text):
    """Add runs to a paragraph, handling **bold**, *italic*, and inline code."""
    # Handle bold (**text**) and italic (*text*)
    # This is a simple parser, not a full markdown parser
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)'
    for part in re.findall(pattern, text):
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Consolas'
        else:
            run = para.add_run(part)

def process_markdown_blocks(content):
    """Yield blocks of (type, lines) from markdown-style content."""
    lines = content.split('\n')
    block = []
    block_type = None
    for line in lines + [None]:  # Add sentinel
        l = line.strip() if line is not None else None
        if l is None or l == '':
            if block:
                yield (block_type, block)
                block = []
                block_type = None
            continue
        # Detect block type
        if re.match(r'^(#+) ', l):
            if block:
                yield (block_type, block)
                block = []
            block_type = 'heading'
            block = [l]
            yield (block_type, block)
            block = []
            block_type = None
        elif re.match(r'^[-*] ', l):
            if block_type != 'ulist':
                if block:
                    yield (block_type, block)
                block = []
                block_type = 'ulist'
            block.append(l)
        elif re.match(r'^\d+\. ', l):
            if block_type != 'olist':
                if block:
                    yield (block_type, block)
                block = []
                block_type = 'olist'
            block.append(l)
        else:
            if block_type != 'para':
                if block:
                    yield (block_type, block)
                block = []
                block_type = 'para'
            block.append(l)
    if block:
        yield (block_type, block)

def add_markdown_block(doc, block_type, lines):
    if block_type == 'heading':
        m = re.match(r'^(#+) (.*)', lines[0])
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), min(level, 4))
    elif block_type == 'ulist':
        for l in lines:
            m = re.match(r'^[-*] (.*)', l)
            if m:
                para = doc.add_paragraph(style='List Bullet')
                add_runs_with_formatting(para, m.group(1))
    elif block_type == 'olist':
        for l in lines:
            m = re.match(r'^\d+\. (.*)', l)
            if m:
                para = doc.add_paragraph(style='List Number')
                add_runs_with_formatting(para, m.group(1))
    elif block_type == 'para':
        # Join lines for a single paragraph
        para = doc.add_paragraph()
        add_runs_with_formatting(para, ' '.join(lines))

def extract_actual_content(raw_content):
    """Extract content after 'Content:' and ignore template/instruction blocks."""
    lines = raw_content.split('\n')
    content_lines = []
    content_started = False
    for line in lines:
        if not content_started:
            if line.strip().lower().startswith('content:'):
                content_started = True
                continue
            # Skip template/instruction lines
            continue
        content_lines.append(line)
    return '\n'.join(content_lines).strip()

def create_proper_docx_file(filename, content, folder):
    """Create a proper Word document with the content"""
    doc = Document()
    
    # Extract title from filename
    title = filename.replace('-', ' ').replace('_', ' ').replace('.docx', '').title()
    
    # Add meta information
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Meta Title: {title}")
    title_run.bold = True
    title_run.font.size = Inches(0.14)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    desc_para = doc.add_paragraph()
    desc_run = desc_para.add_run(f"Meta Description: Comprehensive guide about {title.lower()}. Learn everything you need to know.")
    desc_run.italic = True
    desc_run.font.size = Inches(0.12)
    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add separator
    doc.add_paragraph("=" * 80)
    doc.add_paragraph()
    
    # Extract only the actual content
    actual_content = extract_actual_content(content)
    if not actual_content.strip():
        print(f"WARNING: No main content found after 'Content:' in {filename}")
    # Process blocks as before
    first_line = True
    for block_type, lines in process_markdown_blocks(actual_content):
        # First non-empty block is the main title (Heading 1)
        if first_line and block_type == 'para' and lines:
            doc.add_heading(' '.join(lines), 1)
            first_line = False
            continue
        # Section headers as Heading 2 (if para and all-caps or ends with ':')
        if block_type == 'para' and len(lines) == 1:
            l = lines[0].strip()
            if l.isupper() or (l.endswith(':') and len(l.split()) < 8):
                doc.add_heading(l.rstrip(':'), 2)
                continue
        add_markdown_block(doc, block_type, lines)
    # Add hashtags as a final paragraph if present
    hashtags = [l for l in actual_content.split('\n') if l.strip().startswith('#')]
    if hashtags:
        doc.add_paragraph(' '.join(hashtags))
    
    # Save the document
    filepath = os.path.join(folder, filename)
    doc.save(filepath)
    return filepath

def main():
    """Main function to create all documents"""
    print("Creating all blog posts and service documents as proper Word files...")
    print("=" * 70)
    
    # Create folders if they don't exist
    for folder in ['blog-posts', 'service']:
        if not os.path.exists(folder):
            os.makedirs(folder)
            print(f"Created {folder} folder")
    
    # Create service blog posts
    print("\nCreating service blog posts...")
    service_count = 0
    for service in ALL_SERVICES:
        filename = f"{service}-blog-post.docx"
        content = create_service_blog_content(service)
        filepath = create_proper_docx_file(filename, content, 'service')
        service_count += 1
        print(f"✓ Created: {filepath}")
    
    # Create general blog posts
    print("\nCreating general blog posts...")
    blog_count = 0
    for topic in BLOG_POSTS:
        filename = f"{topic}-blog-post.docx"
        content = create_general_blog_content(topic)
        filepath = create_proper_docx_file(filename, content, 'blog-posts')
        blog_count += 1
        print(f"✓ Created: {filepath}")
    
    print("\n" + "=" * 70)
    print("Document creation complete!")
    print(f"Service blog posts created: {service_count}")
    print(f"General blog posts created: {blog_count}")
    print(f"Total documents created: {service_count + blog_count}")
    print("\nAll files are now proper Word documents (.docx) that can be opened in Microsoft Word!")

if __name__ == "__main__":
    main() 